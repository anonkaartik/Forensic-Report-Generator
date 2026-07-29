from pathlib import Path

from docx import Document

from app.reports.report_builder import ReportBuilder


class DOCXExport:

    @staticmethod
    def export(report, output_path):

        data = ReportBuilder.build(report)

        Path(output_path).parent.mkdir(
            parents=True,
            exist_ok=True
        )

        document = Document()

        document.add_heading(
            "Digital Forensic Investigation Report",
            level=1
        )

        document.add_heading("Case Information", level=2)

        document.add_paragraph(f"Case Name: {data['case_name']}")
        document.add_paragraph(f"Case Number: {data['case_number']}")
        document.add_paragraph(f"Examiner: {data['examiner']}")
        document.add_paragraph(f"Organization: {data['organization']}")
        document.add_paragraph(f"Report Date: {data['report_date']}")

        document.add_heading("Executive Summary", level=2)
        document.add_paragraph(data["executive_summary"])

        document.add_heading("Methodology", level=2)
        document.add_paragraph(data["methodology"])

        document.add_heading("Findings", level=2)

        for finding in data["findings"]:
            document.add_paragraph(
                finding,
                style="List Bullet"
            )

        document.add_heading("Tool Versions", level=2)

        table = document.add_table(rows=1, cols=2)

        table.style = "Table Grid"

        header = table.rows[0].cells
        header[0].text = "Tool"
        header[1].text = "Version"

        for tool in data["tool_versions"]:
            row = table.add_row().cells
            row[0].text = tool["tool_name"]
            row[1].text = tool["version"]

        document.add_heading("Hash Log", level=2)

        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        header = table.rows[0].cells
        header[0].text = "File"
        header[1].text = "Algorithm"
        header[2].text = "Hash"

        for item in data["hash_log"]:
            row = table.add_row().cells
            row[0].text = item["file_name"]
            row[1].text = item["algorithm"]
            row[2].text = item["hash_value"]

        document.save(output_path)

        return output_path